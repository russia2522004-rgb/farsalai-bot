import os
import json
import copy
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from database import get_equipment_by_model, get_equipment_blocks

OUTPUT_DIR = 'output'
TEMPLATE_PATH = 'template/kp_template.docx'
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs('template', exist_ok=True)


def _replace_text_in_runs(paragraph, placeholder, value):
    full_text = ''.join([run.text for run in paragraph.runs])
    if placeholder not in full_text:
        return False
    new_full = full_text.replace(placeholder, str(value) if value else '')
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for run in paragraph.runs[1:]:
            run.text = ''
    return True


def _replace_in_document(doc, replacements: dict):
    for paragraph in doc.paragraphs:
        for k, v in replacements.items():
            _replace_text_in_runs(paragraph, k, v)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for k, v in replacements.items():
                        _replace_text_in_runs(paragraph, k, v)


def _find_content_placeholder(doc):
    for para in doc.paragraphs:
        if '{{CONTENT}}' in ''.join(r.text for r in para.runs):
            return para
    return None


def _add_equipment_header(doc, insert_after_elem, name: str):
    """Заголовок оборудования — чёрный фон, белый текст, 18pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '000000')
    pPr.append(shd)
    run = p.add_run(name)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    insert_after_elem.addnext(p._element)
    return p._element


def _add_section_title(doc, insert_after_elem, title: str, number: int = 0):
    """Заголовок раздела — серый фон, белый текст, keepNext"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '595959')
    pPr.append(shd)
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)
    display_title = f"{number}. {title}" if number else title
    run = p.add_run(display_title)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    insert_after_elem.addnext(p._element)
    return p._element


def _add_horizontal_line(doc, insert_after_elem, keep_next=False):
    """Горизонтальная линия"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    if keep_next:
        kn = OxmlElement('w:keepNext')
        pPr.append(kn)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    insert_after_elem.addnext(p._element)
    return p._element


def _set_keep_next(elem):
    pPr = elem.find(f'{{{NS}}}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        elem.insert(0, pPr)
    kn = OxmlElement('w:keepNext')
    pPr.append(kn)


def _set_cant_split_first_rows(tbl_elem, rows=2):
    """Запрещает разрывать первые строки таблицы + повторяет шапку"""
    tr_list = tbl_elem.findall(f'{{{NS}}}tr')
    for i, tr in enumerate(tr_list[:rows]):
        trPr = tr.find(f'{{{NS}}}trPr')
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        if i == 0:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)


def _get_first_content_type(xml_content: str) -> str:
    """Определяет тип первого значимого элемента в блоке"""
    try:
        from lxml import etree
        DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        block = etree.fromstring(xml_content)
        for child in block:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tbl':
                return 'table'
            if tag == 'p':
                if list(child.iter(f'{{{DRAW_NS}}}inline')) or list(child.iter(f'{{{DRAW_NS}}}anchor')):
                    return 'image'
                texts = [t.text or '' for t in child.iter(f'{{{NS}}}t')]
                if ''.join(texts).strip():
                    return 'text'
    except Exception:
        pass
    return 'text'


def _add_images_to_doc(doc, images_b64: list) -> dict:
    """
    Добавляет картинки из base64 в документ как relationships.
    Возвращает маппинг порядкового номера → новый rId.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    rid_map = {}
    for i, b64_str in enumerate(images_b64):
        if not b64_str or not b64_str.startswith('data:image/'):
            continue
        try:
            header, data = b64_str.split(',', 1)
            mime = header.split('/')[1].split(';')[0]
            ext = 'jpeg' if mime == 'jpeg' else mime
            img_data = base64.b64decode(data)

            img_part = Part(
                PackURI(f'/word/media/block_img_{i}_{id(doc)}.{ext}'),
                f'image/{mime}',
                img_data
            )
            new_rid = doc.part.relate_to(
                img_part,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
            )
            rid_map[i] = new_rid
        except Exception as e:
            print(f"Ошибка добавления картинки {i}: {e}")

    return rid_map


def _update_rids_in_xml(xml_content: str, rid_map: dict) -> str:
    """
    Обновляет rId ссылки в XML блока.
    rid_map: {порядковый_номер_картинки → новый_rId}
    """
    try:
        from lxml import etree
        A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

        block = etree.fromstring(xml_content)
        blips = list(block.iter(f'{{{A_NS}}}blip'))

        for i, blip in enumerate(blips):
            if i in rid_map:
                blip.set(f'{{{REL_NS}}}embed', rid_map[i])

        return etree.tostring(block, encoding='unicode')
    except Exception as e:
        print(f"Ошибка обновления rId: {e}")
        return xml_content


def _insert_xml_block(doc, insert_after_elem, xml_content: str, rid_map: dict = None):
    """Вставляет XML блок с правильными rId для картинок"""
    try:
        from lxml import etree

        # Обновляем rId если есть маппинг
        if rid_map:
            xml_content = _update_rids_in_xml(xml_content, rid_map)

        block = etree.fromstring(xml_content)
        children = list(block)
        if not children:
            return False

        first_type = _get_first_content_type(xml_content)

        # Вставляем все элементы в обратном порядке
        for child in reversed(children):
            child_copy = copy.deepcopy(child)
            insert_after_elem.addnext(child_copy)

        # Настройка разрывов страниц
        parent = insert_after_elem.getparent()
        all_elems = list(parent)
        start_idx = all_elems.index(insert_after_elem) + 1
        if start_idx < len(all_elems):
            first_inserted = all_elems[start_idx]
            tag = first_inserted.tag.split('}')[-1] if '}' in first_inserted.tag else first_inserted.tag

            if first_type == 'table':
                _set_cant_split_first_rows(first_inserted, rows=2)
                anchor = OxmlElement('w:p')
                anchorPr = OxmlElement('w:pPr')
                kn = OxmlElement('w:keepNext')
                anchorPr.append(kn)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                anchorPr.append(spacing)
                anchor.append(anchorPr)
                first_inserted.addprevious(anchor)
            elif first_type == 'image':
                if tag == 'p':
                    _set_keep_next(first_inserted)

        return True
    except Exception as e:
        print(f"Ошибка вставки XML блока: {e}")
        return False


def _add_conditions_block(doc, insert_after_elem, item: dict, eq: dict):
    """Блок условий с линиями сверху и снизу"""
    production_time = item.get('production_time') or (eq.get('production_time') if eq else None) or '25-30 дней'
    packaging = item.get('packaging') or (eq.get('packaging') if eq else None) or 'экспортная деревянная тара (ящик)'
    delivery = item.get('delivery') or (eq.get('delivery') if eq else None) or 'до завода покупателя'
    payment_terms = item.get('payment_terms') or (eq.get('payment_terms') if eq else None) or '50% – предоплата, 50% – по факту поставки'
    unit_price = item.get('unit_price', 0)
    currency = item.get('currency', 'ЮАНЕЙ')

    conditions = [
        ('Цена с НДС с доставкой ' + delivery + ' за 1 штуку:', f'{unit_price:,.0f} {currency}.'),
        ('Условия оплаты:', payment_terms + '.'),
        ('Упаковка:', packaging + '.'),
        ('Сроки изготовления:', production_time + '.'),
    ]

    # Нижняя линия
    _add_horizontal_line(doc, insert_after_elem)

    for label, value in conditions:
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        kl = OxmlElement('w:keepLines')
        pPr.append(kl)
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.name = 'Arial'
        run_value = p.add_run(value)
        run_value.font.size = Pt(10)
        run_value.font.name = 'Arial'
        insert_after_elem.addnext(p._element)

    # Верхняя линия — держится с условиями
    _add_horizontal_line(doc, insert_after_elem, keep_next=True)

    return insert_after_elem


def _add_summary_table(doc, insert_after_elem, items: list):
    """Итоговая таблица для нескольких позиций"""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['Оборудование', 'Кол-во', 'Цена за ед.', 'Сумма']
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Arial'

    total = 0
    currency = ''
    for item in items:
        row = table.add_row().cells
        row[0].text = item.get('name', item.get('model', ''))
        qty = item.get('quantity', 1)
        price = item.get('unit_price', 0)
        currency = item.get('currency', 'ЮАНЕЙ')
        subtotal = price * qty
        total += subtotal
        row[1].text = str(qty)
        row[2].text = f"{price:,.0f} {currency}"
        row[3].text = f"{subtotal:,.0f} {currency}"
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'

    total_row = table.add_row().cells
    total_row[0].text = 'ИТОГО'
    total_row[3].text = f"{total:,.0f} {currency}"
    for cell in total_row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Arial'

    insert_after_elem.addnext(table._tbl)
    title_p = doc.add_paragraph()
    title_r = title_p.add_run('Итоговая стоимость')
    title_r.bold = True
    title_r.font.size = Pt(11)
    title_r.font.name = 'Arial'
    insert_after_elem.addnext(title_p._element)


def _download_photo(photo_path: str, local_path: str) -> bool:
    """Скачивает фото с Яндекс Диска"""
    try:
        import requests
        token = os.getenv('YANDEX_DISK_TOKEN')
        headers = {'Authorization': f'OAuth {token}'}
        r = requests.get('https://cloud-api.yandex.net/v1/disk/resources/download',
                         headers=headers, params={'path': photo_path}, timeout=30)
        if r.status_code == 200:
            download_url = r.json().get('href')
            if download_url:
                img_r = requests.get(download_url, timeout=30)
                with open(local_path, 'wb') as f:
                    f.write(img_r.content)
                return True
        else:
            print(f"Ошибка получения ссылки фото {photo_path}: {r.status_code}")
    except Exception as e:
        print(f"Ошибка скачивания фото: {e}")
    return False


def _apply_numbering_xml(doc, numbering_xml: str):
    """Копирует numbering.xml из оригинального документа в текущий"""
    if not numbering_xml:
        return
    try:
        from lxml import etree
        from docx.oxml.ns import nsmap
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        src_root = etree.fromstring(numbering_xml.encode('utf-8'))

        # Проверяем есть ли уже numbering part в документе
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            # Создаём новый numbering part
            num_part = Part(
                PackURI('/word/numbering.xml'),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
                numbering_xml.encode('utf-8')
            )
            doc.part.relate_to(
                num_part,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
            )
        else:
            # Мёрджим — добавляем недостающие abstractNum и num
            dst_root = numbering_part._element
            NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            existing_abstract = {int(e.get(f'{{{NS_W}}}abstractNumId', 0))
                                 for e in dst_root.findall(f'{{{NS_W}}}abstractNum')}
            existing_num = {int(e.get(f'{{{NS_W}}}numId', 0))
                           for e in dst_root.findall(f'{{{NS_W}}}num')}

            for elem in src_root.findall(f'{{{NS_W}}}abstractNum'):
                aid = int(elem.get(f'{{{NS_W}}}abstractNumId', -1))
                if aid not in existing_abstract:
                    dst_root.append(copy.deepcopy(elem))

            for elem in src_root.findall(f'{{{NS_W}}}num'):
                nid = int(elem.get(f'{{{NS_W}}}numId', -1))
                if nid not in existing_num:
                    dst_root.append(copy.deepcopy(elem))

    except Exception as e:
        print(f"Ошибка применения numbering.xml: {e}")


def generate_kp_document(kp_data: dict, manager_name: str) -> tuple[str, str]:
    """Генерирует КП из блоков оборудования"""
    kp_number = kp_data.get('kp_number', 'KP-001')
    kp_date = kp_data.get('kp_date', datetime.now().strftime('%d.%m.%Y'))
    items = kp_data.get('items', [])

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    # Заменяем шапку
    _replace_in_document(doc, {
        '{{DATE}}': kp_date,
        '{{KP_NUMBER}}': kp_number,
    })

    # keepLines на подпись
    for para in doc.paragraphs:
        if 'уважением' in para.text or 'Лавришко' in para.text:
            pPr = para._element.get_or_add_pPr()
            kl = OxmlElement('w:keepLines')
            pPr.append(kl)

    # Находим {{CONTENT}}
    content_para = _find_content_placeholder(doc)
    if content_para is None:
        raise ValueError("Плейсхолдер {{CONTENT}} не найден в шаблоне")

    for run in content_para.runs:
        run.text = ''

    insert_after = content_para._element

    # Итоговая таблица если несколько позиций
    if len(items) > 1:
        _add_summary_table(doc, insert_after, items)

    # Обрабатываем позиции в обратном порядке
    for item in reversed(items):
        model = item.get('model', '')
        eq = get_equipment_by_model(model)
        blocks = get_equipment_blocks(eq['id']) if eq else []

        # Условия позиции
        _add_conditions_block(doc, insert_after, item, eq)

        # Применяем numbering.xml из оригинального документа
        if eq and eq.get('numbering_xml'):
            _apply_numbering_xml(doc, eq['numbering_xml'])

        # Блоки из библиотеки
        total_blocks = len(blocks)
        for idx, block in enumerate(reversed(blocks)):
            block_title = block.get('block_title', '')
            xml_content = block.get('xml_content', '') or block.get('xml', '')
            block_number = total_blocks - idx

            # Добавляем картинки из base64 в документ и получаем маппинг rId
            rid_map = {}
            images_b64 = json.loads(block.get('images_base64', '[]')) if isinstance(block.get('images_base64'), str) else (block.get('images_base64') or [])
            if images_b64:
                rid_map = _add_images_to_doc(doc, images_b64)

            if xml_content:
                _insert_xml_block(doc, insert_after, xml_content, rid_map if rid_map else None)

            if block_title:
                _add_section_title(doc, insert_after, block_title, number=block_number)

        # Фото оборудования
        if eq and eq.get('photo_path'):
            photo_path = eq['photo_path']
            photo_local = f'temp_photo_{kp_number}_{model}.jpg'
            if _download_photo(photo_path, photo_local):
                try:
                    try:
                        from PIL import Image as PILImage
                        with PILImage.open(photo_local) as pil_img:
                            orig_w, orig_h = pil_img.size
                        w_cm = 14.0
                        h_cm = orig_h / orig_w * w_cm
                        if h_cm > 9.0:
                            h_cm = 9.0
                            w_cm = orig_w / orig_h * h_cm
                    except Exception:
                        w_cm = 12.0

                    space_p = doc.add_paragraph()
                    insert_after.addnext(space_p._element)

                    note_p = doc.add_paragraph()
                    note_r = note_p.add_run(
                        '* Фото для справки. Реальные фотографии будут предоставлены после завершения производства.')
                    note_r.font.size = Pt(9)
                    note_r.italic = True
                    note_r.font.name = 'Arial'
                    insert_after.addnext(note_p._element)

                    photo_p = doc.add_paragraph()
                    photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pPr = photo_p._element.get_or_add_pPr()
                    kn = OxmlElement('w:keepNext')
                    pPr.append(kn)
                    run_photo = photo_p.add_run()
                    run_photo.add_picture(photo_local, width=Cm(w_cm))
                    insert_after.addnext(photo_p._element)
                except Exception as e:
                    print(f"Ошибка вставки фото: {e}")
                finally:
                    if os.path.exists(photo_local):
                        os.remove(photo_local)

        # Заголовок оборудования
        name = eq['name'] if eq else item.get('name', model)
        _add_equipment_header(doc, insert_after, name)

        if len(items) > 1:
            sep = doc.add_paragraph()
            insert_after.addnext(sep._element)

    docx_path = os.path.join(OUTPUT_DIR, f'КП_{kp_number}.docx')
    doc.save(docx_path)
    pdf_path = _convert_to_pdf(kp_data, kp_number)
    return docx_path, pdf_path


def _convert_to_pdf(kp_data: dict, kp_number: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdf_path = os.path.join(OUTPUT_DIR, f'КП_{kp_number}.pdf')

    font_name, font_bold = 'Helvetica', 'Helvetica-Bold'
    for regular, bold in [
        ('fonts/DejaVuSans.ttf', 'fonts/DejaVuSans-Bold.ttf'),
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
    ]:
        if os.path.exists(regular):
            try:
                pdfmetrics.registerFont(TTFont('CyrFont', regular))
                pdfmetrics.registerFont(TTFont('CyrFont-Bold', bold if os.path.exists(bold) else regular))
                font_name, font_bold = 'CyrFont', 'CyrFont-Bold'
                break
            except Exception:
                pass

    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                                leftMargin=2*cm, rightMargin=1.5*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

    n = ParagraphStyle('n', fontName=font_name, fontSize=10, leading=14)
    b = ParagraphStyle('b', fontName=font_bold, fontSize=10, leading=14)
    t = ParagraphStyle('t', fontName=font_bold, fontSize=14, leading=18, alignment=1)
    s = ParagraphStyle('s', fontName=font_bold, fontSize=12, leading=16,
                       backColor=colors.black, textColor=colors.white, alignment=1)
    r = ParagraphStyle('r', fontName=font_name, fontSize=10, leading=14, alignment=2)

    story = []
    kp_date = kp_data.get('kp_date', datetime.now().strftime('%d.%m.%Y'))
    kp_num = kp_data.get('kp_number', kp_number)
    items = kp_data.get('items', [])

    story.append(Paragraph(f'от {kp_date}    №    {kp_num}', r))
    story.append(Paragraph('г. Таганрог', r))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', t))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph('ООО «Фарсал» предлагает к поставке',
                           ParagraphStyle('sub', fontName=font_bold, fontSize=12, alignment=1)))
    story.append(Spacer(1, 0.3*cm))

    for item in items:
        model = item.get('model', '')
        eq = get_equipment_by_model(model)
        name = eq['name'] if eq else item.get('name', model)
        story.append(Paragraph(name, s))
        story.append(Spacer(1, 0.3*cm))

        if eq and eq.get('specs'):
            try:
                specs = json.loads(eq['specs']) if isinstance(eq['specs'], str) else eq['specs']
                if specs:
                    story.append(Paragraph('Технические характеристики', b))
                    story.append(Spacer(1, 0.2*cm))
                    td = [['Характеристика', 'Значение']]
                    for sp in specs:
                        td.append([sp.get('name', ''), str(sp.get('value', ''))])
                    tbl = Table(td, colWidths=[9*cm, 8*cm])
                    tbl.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, 0), font_bold),
                        ('FONTNAME', (0, 1), (-1, -1), font_name),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 0.4*cm))
            except Exception:
                pass

        production_time = item.get('production_time') or (eq.get('production_time') if eq else None) or '25-30 дней'
        packaging = item.get('packaging') or (eq.get('packaging') if eq else None) or 'экспортная деревянная тара (ящик)'
        delivery = item.get('delivery') or (eq.get('delivery') if eq else None) or 'до завода покупателя'
        payment_terms = item.get('payment_terms') or (eq.get('payment_terms') if eq else None) or '50% предоплата, 50% по факту'
        unit_price = item.get('unit_price', 0)
        currency = item.get('currency', 'ЮАНЕЙ')

        story.append(Paragraph(f'Сроки изготовления: {production_time}.', n))
        story.append(Paragraph(f'Упаковка: {packaging}.', n))
        story.append(Paragraph(f'Условия оплаты: {payment_terms}.', n))
        story.append(Paragraph(f'Цена с НДС с доставкой {delivery} за 1 шт.: {unit_price:,.0f} {currency}.', b))
        story.append(Spacer(1, 0.5*cm))

    if len(items) > 1:
        story.append(Paragraph('Итоговая стоимость', b))
        td = [['Оборудование', 'Кол-во', 'Цена/шт', 'Сумма']]
        total = 0
        curr = ''
        for item in items:
            qty = item.get('quantity', 1)
            price = item.get('unit_price', 0)
            curr = item.get('currency', 'ЮАНЕЙ')
            sub = price * qty
            total += sub
            td.append([item.get('name', item.get('model', '')), str(qty),
                       f"{price:,.0f} {curr}", f"{sub:,.0f} {curr}"])
        td.append(['ИТОГО', '', '', f"{total:,.0f} {curr}"])
        tbl = Table(td, colWidths=[8*cm, 2*cm, 4*cm, 4*cm])
        tbl.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), font_bold),
            ('FONTNAME', (0, -1), (-1, -1), font_bold),
            ('FONTNAME', (0, 1), (-1, -2), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.5*cm))

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph('С уважением,', n))
    story.append(Paragraph('директор ООО «Фарсал»,', n))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph('МП     _______________       А. Ю. Лавришко', n))

    pdf_doc.build(story)
    return pdf_path


def cleanup_temp_files(docx_path: str, pdf_path: str):
    for path in [docx_path, pdf_path]:
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass
