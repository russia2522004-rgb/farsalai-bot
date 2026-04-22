import os
import json
import logging
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    filters, ContextTypes, ConversationHandler
)
from openai import OpenAI

from database import (
    init_db, add_equipment, get_all_equipment, search_equipment,
    get_equipment_by_model, delete_equipment, update_equipment,
    save_kp, update_kp, get_kp_by_number, search_kp, get_recent_kp,
    generate_kp_number, save_equipment_blocks, get_equipment_blocks
)
from claude_agent import chat_with_claude, process_edit, extract_equipment_from_doc, extract_all_equipment_from_doc, compare_equipment, resolve_equipment_conflict
from document_generator import generate_kp_document, cleanup_temp_files
from storage import upload_kp_files, add_kp_to_sheets, update_kp_in_sheets, ensure_headers, upload_equipment_photo

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

MAIN_MENU, CREATING_KP, EDITING_KP, ADDING_EQUIPMENT, SEARCHING_KP = range(5)

ALLOWED_IDS = set(map(int, os.getenv('ALLOWED_USER_IDS', '').split(',')))
ADMIN_ID = int(os.getenv('ALLOWED_USER_IDS', '0').split(',')[0])
LOG_CHANNEL_ID = os.getenv('LOG_CHANNEL_ID')

# Имена менеджеров
MANAGER_NAMES = {
    177592975: 'ЛавАлеСер',
    922595157: 'ЛавАлеЮр',
}

openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
user_sessions = {}


def is_allowed(user_id: int) -> bool:
    return user_id in ALLOWED_IDS


def get_manager_name(user_id: int, fallback: str) -> str:
    return MANAGER_NAMES.get(user_id, fallback)


def get_session(user_id: int) -> dict:
    if user_id not in user_sessions:
        user_sessions[user_id] = {
            'state': MAIN_MENU,
            'conversation': [],
            'kp_data': None,
            'kp_number': None,
            'sheets_row': None,
        }
    return user_sessions[user_id]


async def send_log(context, message: str):
    """Отправляет сообщение в канал логов"""
    if not LOG_CHANNEL_ID:
        return
    try:
        await context.bot.send_message(
            chat_id=LOG_CHANNEL_ID,
            text=message,
            parse_mode='HTML'
        )
    except Exception as e:
        logger.error(f"Ошибка отправки лога: {e}")


async def log_bot_response(context, manager_name: str, text: str):
    """Логирует ответ бота менеджеру"""
    if not LOG_CHANNEL_ID:
        return
    # Обрезаем длинные ответы
    short_text = text[:500] + '...' if len(text) > 500 else text
    # Убираем markdown разметку для лога
    clean = short_text.replace('*', '').replace('_', '').replace('`', '')
    await send_log(context,
        f"🤖 <b>Ответ бота</b>\n"
        f"👤 Менеджеру: {manager_name}\n"
        f"💬 {clean}"
    )


async def log_files(context, manager_name: str, kp_number: str, docx_path: str, pdf_path: str):
    """Пересылает файлы КП в канал логов"""
    if not LOG_CHANNEL_ID:
        return
    try:
        await context.bot.send_message(
            chat_id=LOG_CHANNEL_ID,
            text=f"📎 <b>Файлы КП №{kp_number}</b>\n👤 Менеджер: {manager_name}",
            parse_mode='HTML'
        )
        if os.path.exists(docx_path):
            with open(docx_path, 'rb') as f:
                await context.bot.send_document(
                    chat_id=LOG_CHANNEL_ID,
                    document=f,
                    filename=f'КП_{kp_number}.docx'
                )
        if pdf_path and pdf_path.endswith('.pdf') and os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                await context.bot.send_document(
                    chat_id=LOG_CHANNEL_ID,
                    document=f,
                    filename=f'КП_{kp_number}.pdf'
                )
    except Exception as e:
        logger.error(f"Ошибка отправки файлов в лог: {e}")


async def transcribe_voice(file_path: str) -> str:
    with open(file_path, 'rb') as audio:
        transcript = openai_client.audio.transcriptions.create(
            model='whisper-1',
            file=audio,
            language='ru'
        )
    return transcript.text


async def handle_voice_or_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> str:
    user = update.effective_user
    manager_name = get_manager_name(user.id, user.full_name)

    if update.message.voice:
        voice = update.message.voice
        file = await context.bot.get_file(voice.file_id)
        voice_path = f'temp_voice_{user.id}.ogg'
        await file.download_to_drive(voice_path)
        await update.message.reply_text('🎙 Распознаю голос...')
        text = await transcribe_voice(voice_path)
        os.remove(voice_path)
        await update.message.reply_text(f'📝 Распознано: _{text}_', parse_mode='Markdown')

        # Лог голосового
        await send_log(context,
            f"🎙 <b>Голосовое сообщение</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"📝 Транскрипция: {text}"
        )
        return text

    text = update.message.text
    # Лог текстового (только если не кнопка меню)
    menu_buttons = ['📄 Создать КП', '📚 Библиотека оборудования', '🔍 Найти КП',
                    '📋 Последние КП', '✅ Готово, сохранить', '✏️ Внести правки', '❌ Отменить']
    if text and text not in menu_buttons and not text.startswith('/'):
        await send_log(context,
            f"💬 <b>Текстовое сообщение</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"📝 Текст: {text}"
        )
    return text


# ─── Команды ─────────────────────────────────────────────────────────────────

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        await update.message.reply_text('⛔ У вас нет доступа к этому боту.')
        return

    session = get_session(user_id)
    session['state'] = MAIN_MENU

    keyboard = [
        ['📄 Создать КП'],
        ['📚 Библиотека оборудования', '🔍 Найти КП'],
        ['📋 Последние КП'],
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text(
        '👋 Добро пожаловать в *ФарсалИИ*!\n\nВыберите действие:',
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_allowed(update.effective_user.id):
        return
    await update.message.reply_text(
        '*Доступные команды:*\n\n'
        '/start — главное меню\n'
        '/newkp — создать новое КП\n'
        '/cancel — отменить текущее действие\n'
        '/equipment — библиотека оборудования\n'
        '/history — последние КП\n'
        '/find — найти КП',
        parse_mode='Markdown'
    )


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        return
    session = get_session(user_id)
    session['state'] = MAIN_MENU
    session['conversation'] = []
    session['kp_data'] = None

    keyboard = [
        ['📄 Создать КП'],
        ['📚 Библиотека оборудования', '🔍 Найти КП'],
        ['📋 Последние КП'],
    ]
    await update.message.reply_text(
        '❌ Действие отменено.',
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )


# ─── Создание КП ─────────────────────────────────────────────────────────────

async def start_kp_creation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    session = get_session(user_id)
    session['state'] = CREATING_KP
    session['conversation'] = []
    session['kp_data'] = None

    manager_name = get_manager_name(user_id, update.effective_user.full_name)
    await send_log(context,
        f"📄 <b>Начало создания КП</b>\n"
        f"👤 Менеджер: {manager_name}"
    )

    await update.message.reply_text(
        '📄 *Создание нового КП*\n\n'
        'Опишите что нужно включить в коммерческое предложение.\n'
        'Можно голосом или текстом.',
        parse_mode='Markdown',
        reply_markup=ReplyKeyboardRemove()
    )


async def process_kp_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    session = get_session(user_id)

    text = await handle_voice_or_text(update, context)
    if not text:
        return

    await context.bot.send_chat_action(update.effective_chat.id, 'typing')
    response, kp_data = chat_with_claude(session['conversation'], text)

    if kp_data:
        session['kp_data'] = kp_data
        await update.message.reply_text('✅ *Данные собраны!* Генерирую документ...', parse_mode='Markdown')

        # Лог собранных данных
        manager_name = get_manager_name(user_id, update.effective_user.full_name)
        items_str = ', '.join([f"{i.get('name', i.get('model'))} x{i.get('quantity')} = {i.get('unit_price')} {i.get('currency', '')}" for i in kp_data.get('items', [])])
        await send_log(context,
            f"✅ <b>Данные КП собраны</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"🏢 Клиент: {kp_data.get('client', '—')}\n"
            f"📦 Позиции: {items_str}\n"
            f"💰 Итого: {kp_data.get('total_price')} {kp_data.get('currency', '')}\n"
            f"💳 Оплата: {kp_data.get('payment_terms', '—')}\n"
            f"⏱ Срок: {kp_data.get('production_time', '—')}"
        )

        await generate_and_send_kp(update, context, session, kp_data)
    else:
        clean_response = response.split('```json')[0].strip() if '```json' in response else response
        await update.message.reply_text(clean_response)
        manager_name = get_manager_name(user_id, update.effective_user.full_name)
        await log_bot_response(context, manager_name, clean_response)


async def generate_and_send_kp(update, context, session, kp_data):
    user_id = update.effective_user.id
    manager_name = get_manager_name(user_id, update.effective_user.full_name)

    models = [item.get('model', '') for item in kp_data.get('items', [])]
    kp_number = generate_kp_number(models)
    kp_data['kp_number'] = kp_number
    kp_data['kp_date'] = datetime.now().strftime('%d.%m.%Y')
    session['kp_number'] = kp_number

    await context.bot.send_chat_action(update.effective_chat.id, 'upload_document')

    try:
        docx_path, pdf_path = generate_kp_document(kp_data, manager_name)

        await update.message.reply_text('☁️ Загружаю на Яндекс Диск...')
        word_url, pdf_url = upload_kp_files(docx_path, pdf_path, kp_number)

        equipment_list = ', '.join([item.get('name', item.get('model')) for item in kp_data.get('items', [])])
        db_data = {
            'kp_number': kp_number,
            'kp_date': kp_data['kp_date'],
            'client': kp_data.get('client'),
            'equipment_list': equipment_list,
            'total_price': kp_data.get('total_price'),
            'currency': kp_data.get('currency', 'ЮАНЕЙ'),
            'payment_terms': kp_data.get('payment_terms'),
            'manager_id': user_id,
            'manager_name': manager_name,
            'yandex_word_url': word_url,
            'yandex_pdf_url': pdf_url,
        }

        sheets_row = add_kp_to_sheets({**db_data}, word_url, pdf_url)
        db_data['sheets_row'] = sheets_row
        session['sheets_row'] = sheets_row
        save_kp(db_data)

        with open(docx_path, 'rb') as f:
            await update.message.reply_document(f, filename=f'КП_{kp_number}.docx')
        if pdf_path.endswith('.pdf') and os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                await update.message.reply_document(f, filename=f'КП_{kp_number}.pdf')

        # Логируем файлы
        await log_files(context, manager_name, kp_number, docx_path, pdf_path)

        cleanup_temp_files(docx_path, pdf_path)

        session['state'] = EDITING_KP

        keyboard = [['✅ Готово, сохранить', '✏️ Внести правки'], ['❌ Отменить']]
        await update.message.reply_text(
            f'📄 *КП №{kp_number} готово!*\n\n'
            f'🔗 [Word]({word_url}) | [PDF]({pdf_url})\n\n'
            'Проверьте документ. Нужны правки?',
            parse_mode='Markdown',
            reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True),
            disable_web_page_preview=True
        )

        # Лог успешного КП
        await send_log(context,
            f"📄 <b>КП сформировано</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"🔢 Номер: {kp_number}\n"
            f"🏢 Клиент: {kp_data.get('client', '—')}\n"
            f"📦 Оборудование: {equipment_list}\n"
            f"🔗 <a href='{word_url}'>Word</a> | <a href='{pdf_url}'>PDF</a>"
        )

    except Exception as e:
        logger.error(f'Ошибка генерации КП: {e}')
        await send_log(context,
            f"❌ <b>Ошибка генерации КП</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"⚠️ Ошибка: {str(e)}"
        )
        await update.message.reply_text(f'❌ Ошибка при генерации КП: {str(e)}')


# ─── Редактирование КП ───────────────────────────────────────────────────────

async def process_edit_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    session = get_session(user_id)

    text = await handle_voice_or_text(update, context)
    if not text:
        return

    if text in ['✅ Готово, сохранить', 'готово', 'сохранить', 'отправляй']:
        await finalize_kp(update, context, session)
        return

    if text in ['❌ Отменить']:
        await cancel(update, context)
        return

    await context.bot.send_chat_action(update.effective_chat.id, 'typing')
    await update.message.reply_text('✏️ Вношу правки...')

    manager_name = get_manager_name(user_id, update.effective_user.full_name)
    await send_log(context,
        f"✏️ <b>Правка КП</b>\n"
        f"👤 Менеджер: {manager_name}\n"
        f"🔢 КП: {session.get('kp_number', '—')}\n"
        f"📝 Правка: {text}"
    )

    response, updated_data = process_edit(session['conversation'], text)

    if updated_data:
        session['kp_data'] = updated_data
        await update.message.reply_text('🔄 Перегенерирую документ...')
        await regenerate_kp(update, context, session, updated_data)
    else:
        await update.message.reply_text(response)
        await log_bot_response(context, manager_name, response)


async def regenerate_kp(update, context, session, kp_data):
    user = update.effective_user
    manager_name = get_manager_name(user.id, user.full_name)
    kp_number = session['kp_number']
    kp_data['kp_number'] = kp_number

    try:
        docx_path, pdf_path = generate_kp_document(kp_data, manager_name)
        word_url, pdf_url = upload_kp_files(docx_path, pdf_path, kp_number)

        if session.get('sheets_row'):
            update_kp_in_sheets(session['sheets_row'], word_url, pdf_url)

        update_kp(kp_number, {
            'yandex_word_url': word_url,
            'yandex_pdf_url': pdf_url,
        })

        with open(docx_path, 'rb') as f:
            await update.message.reply_document(f, filename=f'КП_{kp_number}.docx')
        if pdf_path.endswith('.pdf') and os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                await update.message.reply_document(f, filename=f'КП_{kp_number}.pdf')

        # Логируем файлы
        await log_files(context, manager_name, kp_number, docx_path, pdf_path)

        cleanup_temp_files(docx_path, pdf_path)

        keyboard = [['✅ Готово, сохранить', '✏️ Внести правки'], ['❌ Отменить']]
        await update.message.reply_text(
            f'✅ КП обновлено!\n\n🔗 [Word]({word_url}) | [PDF]({pdf_url})\n\nЕщё правки нужны?',
            parse_mode='Markdown',
            reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True),
            disable_web_page_preview=True
        )

        await send_log(context,
            f"🔄 <b>КП обновлено</b>\n"
            f"👤 Менеджер: {manager_name}\n"
            f"🔢 КП: {kp_number}\n"
            f"🔗 <a href='{word_url}'>Word</a> | <a href='{pdf_url}'>PDF</a>"
        )

    except Exception as e:
        logger.error(f'Ошибка перегенерации: {e}')
        await update.message.reply_text(f'❌ Ошибка: {str(e)}')


async def finalize_kp(update, context, session):
    kp_number = session.get('kp_number', '—')
    manager_name = get_manager_name(update.effective_user.id, update.effective_user.full_name)
    session['state'] = MAIN_MENU
    session['conversation'] = []
    session['kp_data'] = None

    keyboard = [
        ['📄 Создать КП'],
        ['📚 Библиотека оборудования', '🔍 Найти КП'],
        ['📋 Последние КП'],
    ]
    await update.message.reply_text(
        f'✅ *КП №{kp_number} сохранено!*',
        parse_mode='Markdown',
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )

    await send_log(context,
        f"✅ <b>КП финализировано</b>\n"
        f"👤 Менеджер: {manager_name}\n"
        f"🔢 КП: {kp_number}"
    )


# ─── Библиотека оборудования ─────────────────────────────────────────────────

async def equipment_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        return
    equipment = get_all_equipment()
    if not equipment:
        text = '📚 *Библиотека оборудования пуста.*\n\nОтправьте файл КП чтобы добавить оборудование.'
    else:
        lines = ['📚 *Библиотека оборудования:*\n']
        for eq in equipment:
            price = f"{eq['base_price']:,.0f} {eq['currency']}" if eq['base_price'] else 'цена не указана'
            lines.append(f"• *{eq['model']}* — {eq['name']} ({price})")
        text = '\n'.join(lines)
        text += '\n\n_Отправьте файл КП чтобы добавить новое оборудование_'
    await update.message.reply_text(text, parse_mode='Markdown')


async def extract_photos_from_docx(local_path: str) -> list:
    """Извлекает фото из Word документа"""
    photos = []
    try:
        import zipfile
        with zipfile.ZipFile(local_path, 'r') as z:
            media_files = [f for f in z.namelist() if f.startswith('word/media/') and
                          any(f.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp'])]
            for media_file in media_files:
                ext = os.path.splitext(media_file)[1]
                temp_photo = f'temp_photo_{os.getpid()}_{len(photos)}{ext}'
                with z.open(media_file) as src, open(temp_photo, 'wb') as dst:
                    dst.write(src.read())
                # Проверяем размер — пропускаем маленькие картинки (логотипы, иконки)
                size = os.path.getsize(temp_photo)
                if size > 10000:  # больше 10KB — скорее всего фото оборудования
                    photos.append(temp_photo)
                else:
                    os.remove(temp_photo)
    except Exception as e:
        print(f"Ошибка извлечения фото: {e}")
    return photos


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        return

    session = get_session(user_id)
    doc = update.message.document

    if not doc.file_name.endswith(('.docx', '.pdf')):
        await update.message.reply_text('Пожалуйста, отправьте файл в формате .docx или .pdf')
        return

    await update.message.reply_text('📄 Читаю документ...')

    file = await context.bot.get_file(doc.file_id)
    local_path = f'temp_{user_id}_{doc.file_name}'
    await file.download_to_drive(local_path)

    try:
        if doc.file_name.endswith('.docx'):
            from docx import Document as DocxDocument
            d = DocxDocument(local_path)
            paragraphs_text = '\n'.join([p.text for p in d.paragraphs if p.text.strip()])
            tables_text = ''
            for table in d.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text += row_text + '\n'
            doc_text = paragraphs_text + '\n' + tables_text
            photos = await extract_photos_from_docx(local_path)
        else:
            doc_text = 'PDF файл'
            photos = []

        await update.message.reply_text('🤖 Анализирую содержимое...')
        items = extract_all_equipment_from_doc(doc_text, doc_path=local_path)

        if not items:
            await update.message.reply_text('❌ Оборудование не распознано. Возможно это не КП?')
            os.remove(local_path)
            for p in photos:
                if os.path.exists(p):
                    os.remove(p)
            return

        # Распределяем фото по позициям
        for i, item in enumerate(items):
            if i < len(photos):
                item['_photo_path'] = photos[i]
            elif photos:
                item['_photo_path'] = photos[0]
            if 'blocks' not in item:
                item['blocks'] = []

        # Сохраняем очередь для обработки
        session['equipment_queue'] = items
        session['equipment_queue_idx'] = 0
        session['doc_path'] = local_path
        session['state'] = ADDING_EQUIPMENT

        photos_info = f', найдено фото: {len(photos)} шт.' if photos else ', фото не найдено'
        await update.message.reply_text(
            f'📋 Найдено позиций: *{len(items)}*{photos_info}\nОбрабатываю по очереди...',
            parse_mode='Markdown'
        )

        await process_next_equipment(update, context, session)

    except Exception as e:
        logger.error(f'Ошибка обработки документа: {e}')
        await update.message.reply_text(f'❌ Ошибка: {str(e)}')
        if os.path.exists(local_path):
            os.remove(local_path)


async def process_next_equipment(update, context, session):
    """Обрабатывает следующую позицию из очереди"""
    queue = session.get('equipment_queue', [])
    idx = session.get('equipment_queue_idx', 0)

    if idx >= len(queue):
        # Все позиции обработаны
        doc_path = session.get('doc_path')
        if doc_path and os.path.exists(doc_path):
            os.remove(doc_path)
        session['equipment_queue'] = []
        session['state'] = MAIN_MENU

        keyboard = [
            ['📄 Создать КП'],
            ['📚 Библиотека оборудования', '🔍 Найти КП'],
            ['📋 Последние КП'],
        ]
        await update.message.reply_text(
            '✅ Все позиции обработаны!',
            reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
        )
        return

    eq_data = queue[idx]
    model = eq_data.get('model', '')
    existing = get_equipment_by_model(model)

    if not existing:
        # Вариация 1 — новое оборудование
        specs_preview = ''
        if eq_data.get('specs'):
            specs = eq_data['specs'][:3]
            specs_preview = '\n'.join([f"  • {s['name']}: {s['value']}" for s in specs])
            if len(eq_data['specs']) > 3:
                specs_preview += f'\n  • ...ещё {len(eq_data["specs"]) - 3} характеристик'

        text = (
            f'🆕 *Новое оборудование [{idx+1}/{len(queue)}]:*\n\n'
            f'Название: {eq_data.get("name", "—")}\n'
            f'Модель: {model}\n'
            f'Цена: {eq_data.get("base_price", "—")} {eq_data.get("currency", "")}\n'
            f'Срок: {eq_data.get("production_time", "—")}\n'
        )
        if specs_preview:
            text += f'\nХарактеристики:\n{specs_preview}\n'
        text += '\nДобавить в библиотеку?'

        session['pending_equipment'] = eq_data
        keyboard = [['✅ Добавить', '⏭ Пропустить']]
        await update.message.reply_text(
            text,
            parse_mode='Markdown',
            reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
        )

    else:
        # Сравниваем с существующим
        differences = compare_equipment(existing, eq_data)

        if not differences['has_changes']:
            # Вариация 2 — нет изменений, пропускаем автоматически
            await update.message.reply_text(
                f'⏭ *{eq_data.get("name")}* [{idx+1}/{len(queue)}] — уже есть в библиотеке, изменений нет. Пропущено.',
                parse_mode='Markdown'
            )
            session['equipment_queue_idx'] = idx + 1
            await process_next_equipment(update, context, session)

        elif differences['price_changed'] and not differences['specs_changed']:
            # Вариация 4 — только цена изменилась
            pc = differences['price_changed']
            text = (
                f'💰 *{eq_data.get("name")}* [{idx+1}/{len(queue)}]\n\n'
                f'Цена изменилась:\n'
                f'Было: {pc["old"]:,.0f} {pc["currency"]}\n'
                f'Стало: {pc["new"]:,.0f} {pc["currency"]}\n\n'
                f'Обновить цену?'
            )
            session['pending_equipment'] = eq_data
            session['existing_equipment'] = existing
            session['equipment_action'] = 'price_only'
            keyboard = [['✅ Обновить цену', '⏭ Оставить старую']]
            await update.message.reply_text(
                text,
                parse_mode='Markdown',
                reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
            )

        else:
            # Вариация 3 — характеристики отличаются
            diff_lines = []
            for d in differences['specs_changed'][:5]:
                if d['type'] == 'changed':
                    diff_lines.append(f"• {d['name']}: {d['old']} → {d['new']}")
                elif d['type'] == 'added':
                    diff_lines.append(f"• {d['name']}: нет → {d['new']} (новая)")
                elif d['type'] == 'removed':
                    diff_lines.append(f"• {d['name']}: {d['old']} → нет (удалена)")

            if len(differences['specs_changed']) > 5:
                diff_lines.append(f"• ...ещё {len(differences['specs_changed']) - 5} изменений")

            if differences['price_changed']:
                pc = differences['price_changed']
                diff_lines.insert(0, f"• Цена: {pc['old']:,.0f} → {pc['new']:,.0f} {pc['currency']}")

            text = (
                f'⚠️ *{eq_data.get("name")}* [{idx+1}/{len(queue)}]\n\n'
                f'Найдены отличия:\n' +
                '\n'.join(diff_lines) +
                '\n\nЧто сделать? Напишите или надиктуйте:\n'
                '— "обновить всё" — взять все данные из нового файла\n'
                '— "оставить старое" — не менять ничего\n'
                '— или опишите что именно оставить/обновить'
            )
            session['pending_equipment'] = eq_data
            session['existing_equipment'] = existing
            session['equipment_differences'] = differences
            session['equipment_action'] = 'conflict'
            await update.message.reply_text(text, parse_mode='Markdown', reply_markup=ReplyKeyboardRemove())


async def confirm_add_equipment(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    session = get_session(user_id)

    # Получаем текст (голос или текст)
    text = await handle_voice_or_text(update, context)
    if not text:
        return

    action = session.get('equipment_action', 'new')
    idx = session.get('equipment_queue_idx', 0)
    manager_name = get_manager_name(user_id, update.effective_user.full_name)

    if action == 'new':
        # Новое оборудование
        if text in ['✅ Добавить']:
            await update.message.reply_text('⏳ Обрабатываю — загружаю фото и сохраняю блоки...')
            eq_data = session.get('pending_equipment', {})
            photo_path = eq_data.pop('_photo_path', None)
            blocks = eq_data.pop('blocks', [])

            if isinstance(eq_data.get('specs'), list):
                eq_data['specs'] = json.dumps(eq_data['specs'], ensure_ascii=False)

            # Загружаем фото на Яндекс Диск
            if photo_path and os.path.exists(photo_path):
                try:
                    remote_path = upload_equipment_photo(photo_path, eq_data.get('model', 'unknown'))
                    eq_data['photo_path'] = remote_path
                    os.remove(photo_path)
                    await update.message.reply_text('🖼 Фото загружено на Яндекс Диск')
                except Exception as e:
                    logger.error(f"Ошибка загрузки фото: {e}")
                    await update.message.reply_text(f'⚠️ Фото не загружено: {str(e)}')

            eq_id = add_equipment(eq_data)

            # Сохраняем блоки с base64 картинками
            if blocks:
                try:
                    save_equipment_blocks(eq_id, blocks)
                    await update.message.reply_text(f'📦 Сохранено блоков: {len(blocks)}')
                except Exception as e:
                    logger.error(f"Ошибка сохранения блоков: {e}")
            await update.message.reply_text(f'✅ *{eq_data.get("name")}* добавлено!', parse_mode='Markdown')
            await send_log(context,
                f"📚 <b>Добавлено оборудование</b>\n"
                f"👤 {manager_name}\n"
                f"📦 {eq_data.get('name')} (модель: {eq_data.get('model')})"
            )
        else:
            # Удаляем фото если пропускаем
            eq_data = session.get('pending_equipment', {})
            photo_path = eq_data.pop('_photo_path', None)
            if photo_path and os.path.exists(photo_path):
                os.remove(photo_path)
            await update.message.reply_text('⏭ Пропущено.')

        session['equipment_queue_idx'] = idx + 1
        await process_next_equipment(update, context, session)

    elif action == 'price_only':
        if text in ['✅ Обновить цену']:
            eq_data = session.get('pending_equipment', {})
            update_equipment(eq_data.get('model'), {
                'base_price': eq_data.get('base_price'),
                'currency': eq_data.get('currency'),
            })
            await update.message.reply_text('✅ Цена обновлена!', parse_mode='Markdown')
        else:
            await update.message.reply_text('⏭ Цена не изменена.')

        session['equipment_queue_idx'] = idx + 1
        await process_next_equipment(update, context, session)

    elif action == 'conflict':
        # Голосовое/текстовое разрешение конфликта
        existing = session.get('existing_equipment', {})
        new_data = session.get('pending_equipment', {})
        differences = session.get('equipment_differences', {})

        if text.lower() in ['обновить всё', 'обновить все', 'обновить']:
            # Берём все данные из нового файла
            if isinstance(new_data.get('specs'), list):
                new_data['specs'] = json.dumps(new_data['specs'], ensure_ascii=False)
            add_equipment(new_data)
            await update.message.reply_text('✅ Данные полностью обновлены!')
        elif text.lower() in ['оставить старое', 'оставить', 'не менять']:
            await update.message.reply_text('⏭ Оставлены старые данные.')
        else:
            # Claude разрешает конфликт по инструкции
            await update.message.reply_text('🤖 Применяю инструкцию...')
            resolved = resolve_equipment_conflict(existing, new_data, differences, text)
            if isinstance(resolved.get('specs'), list):
                resolved['specs'] = json.dumps(resolved['specs'], ensure_ascii=False)
            add_equipment(resolved)
            await update.message.reply_text('✅ Данные обновлены по вашей инструкции!')

        session['equipment_queue_idx'] = idx + 1
        await process_next_equipment(update, context, session)
    else:
        session['state'] = MAIN_MENU
        await cancel(update, context)


# ─── Просмотр карточки оборудования ─────────────────────────────────────────

async def show_equipment_card(update: Update, context: ContextTypes.DEFAULT_TYPE, model: str):
    """Показывает карточку оборудования"""
    eq = get_equipment_by_model(model)
    if not eq:
        await update.message.reply_text(f'❌ Оборудование "{model}" не найдено в библиотеке.')
        return

    specs_text = ''
    if eq.get('specs'):
        try:
            specs = json.loads(eq['specs']) if isinstance(eq['specs'], str) else eq['specs']
            specs_text = '\n'.join([f"  • {s['name']}: {s['value']}" for s in specs])
        except Exception:
            pass

    price = f"{eq['base_price']:,.0f} {eq['currency']}" if eq.get('base_price') else 'не указана'
    photo_info = '🖼 Фото есть' if eq.get('photo_path') else '📷 Фото нет'

    text = (
        f'📦 *{eq["name"]}*\n\n'
        f'Модель: {eq["model"]}\n'
        f'Цена: {price}\n'
        f'Срок: {eq.get("production_time", "—")}\n'
        f'Упаковка: {eq.get("packaging", "—")}\n'
        f'{photo_info}\n'
    )
    if specs_text:
        text += f'\n*Характеристики:*\n{specs_text}\n'

    # Сохраняем модель для возможного удаления
    session = get_session(update.effective_user.id)
    session['viewing_equipment_model'] = eq['model']

    keyboard = [['🗑 Удалить это оборудование', '◀️ Назад']]
    await update.message.reply_text(
        text,
        parse_mode='Markdown',
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )


async def delete_equipment_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Удаляет оборудование из библиотеки"""
    user_id = update.effective_user.id
    session = get_session(user_id)
    model = session.get('viewing_equipment_model')

    if not model:
        await update.message.reply_text('❌ Не выбрано оборудование для удаления.')
        return

    eq = get_equipment_by_model(model)
    if not eq:
        await update.message.reply_text('❌ Оборудование не найдено.')
        return

    delete_equipment(model)
    manager_name = get_manager_name(user_id, update.effective_user.full_name)

    await send_log(context,
        f"🗑 <b>Удалено оборудование</b>\n"
        f"👤 {manager_name}\n"
        f"📦 {eq.get('name')} (модель: {model})"
    )

    session['viewing_equipment_model'] = None
    keyboard = [
        ['📄 Создать КП'],
        ['📚 Библиотека оборудования', '🔍 Найти КП'],
        ['📋 Последние КП'],
    ]
    await update.message.reply_text(
        f'🗑 *{eq.get("name")}* удалено из библиотеки.',
        parse_mode='Markdown',
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )


# ─── История и поиск КП ──────────────────────────────────────────────────────

async def show_history(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        return

    if user_id == ADMIN_ID:
        kps = get_recent_kp(limit=10)
    else:
        kps = get_recent_kp(manager_id=user_id, limit=10)

    if not kps:
        await update.message.reply_text('📋 История КП пуста.')
        return

    lines = ['📋 *Последние КП:*\n']
    for kp in kps:
        lines.append(
            f"• *{kp['kp_number']}* от {kp['kp_date']}\n"
            f"  {kp.get('client') or '—'} | {kp.get('equipment_list') or '—'}\n"
        )
    await update.message.reply_text('\n'.join(lines), parse_mode='Markdown')


async def find_kp(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        return
    session = get_session(user_id)
    session['state'] = SEARCHING_KP
    await update.message.reply_text(
        '🔍 Введите для поиска:\n• Название клиента\n• Номер КП\n• Название оборудования',
        reply_markup=ReplyKeyboardRemove()
    )


async def process_search(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    session = get_session(user_id)
    query = update.message.text
    results = search_kp(query)
    session['state'] = MAIN_MENU

    if not results:
        await update.message.reply_text('❌ Ничего не найдено.')
    else:
        lines = [f'🔍 *Результаты по "{query}":*\n']
        for kp in results:
            word_link = f'[Word]({kp["yandex_word_url"]})' if kp.get('yandex_word_url') else '—'
            pdf_link = f'[PDF]({kp["yandex_pdf_url"]})' if kp.get('yandex_pdf_url') else '—'
            lines.append(
                f"• *{kp['kp_number']}* от {kp['kp_date']}\n"
                f"  {kp.get('client') or '—'} | {kp.get('equipment_list') or '—'}\n"
                f"  {word_link} | {pdf_link}\n"
            )
        await update.message.reply_text(
            '\n'.join(lines),
            parse_mode='Markdown',
            disable_web_page_preview=True
        )

    keyboard = [
        ['📄 Создать КП'],
        ['📚 Библиотека оборудования', '🔍 Найти КП'],
        ['📋 Последние КП'],
    ]
    await update.message.reply_text('Что дальше?', reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True))


# ─── Главный обработчик ──────────────────────────────────────────────────────

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if not is_allowed(user_id):
        await update.message.reply_text('⛔ У вас нет доступа.')
        return

    session = get_session(user_id)
    text = update.message.text or ''
    state = session['state']

    if text == '📄 Создать КП' or text == '/newkp':
        await start_kp_creation(update, context)
        session['state'] = CREATING_KP
        return
    if text == '📚 Библиотека оборудования' or text == '/equipment':
        await equipment_menu(update, context)
        return
    if text == '🔍 Найти КП' or text == '/find':
        await find_kp(update, context)
        return
    if text == '📋 Последние КП' or text == '/history':
        await show_history(update, context)
        return
    if text == '🗑 Удалить это оборудование':
        await delete_equipment_command(update, context)
        return
    if text == '◀️ Назад':
        await equipment_menu(update, context)
        return
    # Поиск оборудования по названию — "покажи IE-2" или "что есть по F-110"
    if any(kw in text.lower() for kw in ['покажи', 'покажи', 'что есть', 'карточка', 'info']):
        words = text.split()
        for word in words:
            if len(word) > 2 and word not in ['покажи', 'что', 'есть', 'по', 'карточка']:
                await show_equipment_card(update, context, word)
                return

    if state == CREATING_KP:
        await process_kp_message(update, context)
    elif state == EDITING_KP:
        await process_edit_message(update, context)
    elif state == ADDING_EQUIPMENT:
        await confirm_add_equipment(update, context)
    elif state == SEARCHING_KP:
        await process_search(update, context)
    else:
        await update.message.reply_text('Используйте кнопки меню или /start для начала.')


# ─── Запуск ──────────────────────────────────────────────────────────────────

def main():
    init_db()
    try:
        ensure_headers()
    except Exception as e:
        logger.warning(f'Google Sheets недоступен: {e}')

    token = os.getenv('TELEGRAM_BOT_TOKEN')
    app = Application.builder().token(token).build()

    app.add_handler(CommandHandler('start', start))
    app.add_handler(CommandHandler('help', help_command))
    app.add_handler(CommandHandler('cancel', cancel))
    app.add_handler(CommandHandler('newkp', start_kp_creation))
    app.add_handler(CommandHandler('equipment', equipment_menu))
    app.add_handler(CommandHandler('history', show_history))
    app.add_handler(CommandHandler('find', find_kp))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT | filters.VOICE, handle_message))

    logger.info('ФарсалИИ запущен!')
    app.run_polling(drop_pending_updates=True)


if __name__ == '__main__':
    main()
